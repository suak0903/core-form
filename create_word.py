# -*- coding: utf-8 -*-
"""Erstellt core-form_Website-Texte.docx — exakter Text aus den HTML-Quelldateien."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(r"d:\Claude\VS Code\core-form")

HIMBEERE = RGBColor(0xb2, 0x0e, 0x3b)
DUNKEL   = RGBColor(0x0d, 0x0d, 0x0d)
GRAU     = RGBColor(0x2a, 0x2a, 0x2a)

def h1(doc, text):
    doc.add_heading(text, 1)

def h2(doc, text):
    doc.add_heading(text, 2)

def h3(doc, text):
    doc.add_heading(text, 3)

def label(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.color.rgb = HIMBEERE
    r.font.size = Pt(9)

def para(doc, text, bold=False):
    if not text.strip():
        return
    p = doc.add_paragraph()
    r = p.add_run(text.strip())
    r.bold = bold

def divider(doc):
    doc.add_paragraph('─' * 72)

def page(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────

def startseite(doc):
    h1(doc, 'Startseite (index.html)')
    divider(doc)

    # ── Hero ──────────────────────────────────────────────────────────────────
    h2(doc, 'Hero')
    label(doc, 'Pilates & Reformer Studio Essen')
    para(doc, 'Kraft.')
    para(doc, 'Kontrolle.')
    para(doc, 'Bewegung.')
    para(doc, 'Reformer Pilates · Pilates Matte · Barre Workout')
    para(doc, '[Button: Kursangebote entdecken]')
    para(doc, '[Button: Unsere Studios]')

    h3(doc, 'Statistiken')
    para(doc, '2 — Studios in Essen')
    para(doc, '3 — Kursformate')
    para(doc, '35 — Kurse / Woche')

    # ── Manifesto ─────────────────────────────────────────────────────────────
    h2(doc, 'Manifesto-Sektion')
    para(doc, 'Training, das einfach gut tut.')
    para(doc, 'core:form steht für ganzheitliches Bewegen – präzise angeleitet, persönlich begleitet, in kleinen Gruppen. Für Kraft, Stabilität und ein tiefes Wohlbefinden.')

    # ── Angebote ──────────────────────────────────────────────────────────────
    h2(doc, 'Angebote')
    label(doc, 'Was wir anbieten')
    h3(doc, 'Kursangebote')

    karten = [
        ('Studio Südviertel', 'Reformer Pilates',
         'Training an der Reformer-Apparatur in Kleingruppen – präzise angeleitet für maximale Wirkung.',
         'Mehr erfahren →'),
        ('Studio Rüttenscheid', 'Pilates Matte',
         'Achtsam angeleitetes Training in Kleingruppen – zugänglich und tiefgründig zugleich.',
         'Mehr erfahren →'),
        ('Studio Rüttenscheid', 'Barre Workout',
         'Dynamisches Ganzkörpertraining – Pilates trifft Tanz und Kraft. Eleganz trifft Intensität.',
         'Mehr erfahren →'),
        ('Beide Studios', 'Personal Training',
         '1:1 oder 2:1 – individuelles Pilatestraining auf der Matte oder am Reformer.',
         'Anfragen →'),
        ('Studio Südviertel', 'Raumvermietung',
         'Zwei Balanced Body Reformer – ideal für Kolleg:innen und externe Trainer:innen.',
         'Anfragen →'),
    ]
    for tag, titel, desc, link in karten:
        p = doc.add_paragraph()
        p.add_run(f'[{tag}] ').italic = True
        r = p.add_run(titel)
        r.bold = True
        para(doc, desc)
        para(doc, f'[{link}]')

    # ── Feature-Banner ────────────────────────────────────────────────────────
    h2(doc, 'Feature-Banner (Zitat)')
    para(doc, '„Jede Bewegung zählt – wenn sie präzise ist."')
    para(doc, 'Maximal 9 Teilnehmer:innen pro Kurs. Persönliche Anleitung. Echte Wirkung.')

    # ── Studios ───────────────────────────────────────────────────────────────
    h2(doc, 'Studios')
    label(doc, 'Wo wir sind')
    h3(doc, 'Unsere Studios')

    para(doc, 'Studio 01 — Rüttenscheid', bold=True)
    para(doc, 'Pilates Matte & Barre Workout')
    para(doc, '[Button: Info & Buchung]')

    para(doc, 'Studio 02 — Südviertel', bold=True)
    para(doc, 'Reformer Pilates & Raumvermietung')
    para(doc, '[Button: Info & Buchung]')

    # ── Über uns ──────────────────────────────────────────────────────────────
    h2(doc, 'Über uns')
    label(doc, 'Über uns')
    h3(doc, 'Bewegung mit Haltung.')
    para(doc, 'core:form wurde aus der Überzeugung gegründet, dass Pilates mehr ist als ein Workout – es ist eine Praxis. Eine, die Körper und Geist verbindet, präzises Arbeiten mit echter Freude am Bewegen vereint.')
    para(doc, 'Eva, Gründerin und leidenschaftliche Pilates-Trainerin, hat beide Studios in Essen aufgebaut. Mit einem Team erfahrener Trainer:innen bietet core:form ein Umfeld, in dem jede:r das Beste aus sich herausholen kann.')

    h3(doc, 'Werte')
    para(doc, 'Präzision — Jede Bewegung zählt')
    para(doc, 'Wärme — Persönliche Begleitung')
    para(doc, 'Gemeinschaft — Kleine Gruppen, großes Miteinander')

    # ── Team-Karussell ────────────────────────────────────────────────────────
    h2(doc, 'Team-Karussell')
    team = [
        ('Das Team', 'core:form',
         'Ein Kreis erfahrener Pilates- und Bewegungspraktiker:innen, der core:form trägt – präzise angeleitet, persönlich begleitet. Jede:r bringt eigene Schwerpunkte mit, vom Reformer über Barre bis zur klassischen Matte.'),
        ('Eva', 'Inhaberin & Gründerin',
         'Pilates begleitet mich seit über 25 Jahren – seit meinem Tanzstudium, über viele Jahre Unterrichtserfahrung bis hin zur Eröffnung meines ersten Studios 2014. Heute leite ich core:form mit zwei Studios in Essen und bilde Trainer:innen für Matte und Reformer aus. Ich unterrichte vor allem Reformerkurse und fortgeschrittene Mattenstunden.'),
        ('Charlotta', 'Pilates & Barre',
         'Charlotta gehört fast von Anfang an zum core:form-Team. Ihre Pilates-Ausbildung absolvierte sie bei Pilates bodymotion und unterrichtet heute Pilates und Barre Workout. „Mit den Füßen sicher geerdet, mit dem Kopf hoch erhoben und in einer starken Mitte ruhend – das ist mein Pilates-Gefühl."'),
        ('Jelena', 'Reformer & Ausbildung',
         'Pilates begleitet ihren beruflichen Weg seit vielen Jahren. Als ausgebildete Tänzerin bringt Jelena ein ausgeprägtes Körperbewusstsein, Präzision und Bewegungserfahrung mit, die sie durch zahlreiche Aus- und Weiterbildungen im Bereich Pilates und Reformer kontinuierlich vertieft hat. Sie unterrichtet ausschließlich Reformer-Kurse mit großem Fokus auf Technik, individuelle Korrekturen und ein strukturiertes, professionelles Training.'),
        ('Anja', 'Matte, Reformer & Personal',
         'Anja vermittelt fundiert und individuell angepasst die Grundlagen für ein effektives und gesundes Training. Sie unterrichtet auf der Matte, am Reformer und im Personal Training und begleitet Teilnehmer:innen persönlich und achtsam.'),
        ('Nina', 'BASI® Pilates Trainerin',
         'Body & Mind Coach seit rund 20 Jahren, BASI® Pilates Trainerin seit 2012. Nina stärkt vor allem die tiefe Bauchmuskulatur durch abwechslungsreiche, kreative Flows für den ganzen Körper.'),
        ('Pia', 'Pilates Trainerin',
         'Nach ihrer Ausbildung zur Pilates-Trainerin 2025 vereint Pia ihre Erfahrungen als freischaffende Schauspielerin und Tänzerin im Studio. Ihr Fokus: korrekte Form, Anatomie und humorvolle Trainingsmomente.'),
        ('Bille', 'Pilates & Sporttherapie',
         'Ausgebildete Sporttherapeutin und frisch ausgebildete Pilates-Trainerin mit Erfahrung aus der Reha-Arbeit. Ihr Unterricht zeichnet sich durch klare Anleitungen, taktile Korrekturen und eine motivierende Atmosphäre aus.'),
        ('Julia', 'Pilates Trainerin',
         '„Pilates hat mein Leben verändert." Aus der begeisterten Teilnehmerin wurde eine Trainerin mit der Leidenschaft, die positiven Effekte von Pilates weiterzugeben.'),
        ('Cristina', 'Pilates & Barre',
         'Cristina ist Ärztin und seit über sechs Jahren im Pilates und Barre Workout aktiv. Nach ihrer Pilates-Trainerausbildung bei Eva unterrichtet sie auch Barre. Sie legt viel Wert auf kontrollierte Bewegungen und intensive, effektive Muskelarbeit.'),
        ('Rebecca', 'Pilates Trainerin',
         '„Pilates ist für mich mehr als ein Training – es ist eine Einladung, den eigenen Körper neu zu entdecken." Im Unterricht geht es Rebecca um Bewusstsein: bewusste Atmung, bewusste Bewegung, bewusster Umgang mit dem eigenen Körper.'),
    ]
    for name, rolle, bio in team:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        para(doc, rolle)
        para(doc, bio)
        doc.add_paragraph()

    # ── Ausbildung (Teaser) ───────────────────────────────────────────────────
    h2(doc, 'Ausbildung (Teaser-Sektion)')
    label(doc, 'Ausbildung')
    h3(doc, 'Werde Pilates-Trainer:in.')
    para(doc, 'core:form bietet strukturierte Ausbildungsformate für alle, die Pilates professionell unterrichten möchten. Fundiert, praxisnah und von erfahrenen Trainer:innen begleitet.')
    para(doc, 'Premium-Mitglied im Deutschen Pilates Verband')
    para(doc, 'Durchdachte Programme für Matte & Reformer')
    para(doc, 'Erfahrene Ausbilder:innen')
    para(doc, 'Kleine Ausbildungsgruppen')
    para(doc, 'Praxis & Theorie in zwei Studios')
    para(doc, '[Button: Zur Ausbildung]')

    # ── Preise ────────────────────────────────────────────────────────────────
    h2(doc, 'Preise & Pakete')
    label(doc, 'Investition')

    h3(doc, 'Pilates Matte & Barre — ab 15 €')
    label(doc, 'Studio Rüttenscheid')
    para(doc, 'Monatliche Preismodelle')
    para(doc, '1× / Woche — 50 €')
    para(doc, '2× / Woche — 79 €')
    para(doc, 'Probestunde — 15 €')
    para(doc, 'Drop-In — 20 €')
    para(doc, '10er Karten')
    para(doc, '6 Monate Laufzeit — 170 €')
    para(doc, '6 Monate ermäßigt* — 135 €')
    para(doc, '*Rentner:innen, Schüler:innen / Studierende, Alleinerziehende')

    h3(doc, 'Reformer Pilates — ab 25 € [Empfohlen]')
    label(doc, 'Studio Südviertel')
    para(doc, 'Monatliche Preismodelle')
    para(doc, '1× / Woche — 89 €')
    para(doc, '1× / Woche ermäßigt* — 79 €')
    para(doc, '2× / Woche — 160 €')
    para(doc, '2× / Woche ermäßigt* — 144 €')
    para(doc, '10er Karten')
    para(doc, '6 Monate Laufzeit — 250 €')
    para(doc, '6 Monate ermäßigt* — 225 €')
    para(doc, 'Drop-In — 30 €')
    para(doc, 'Probestunde — 25 €')
    para(doc, '*Rentner:innen, Schüler:innen / Studierende, Alleinerziehende, Mitglieder')

    h3(doc, 'Personal Training — ab 40 €')
    label(doc, 'Beide Studios')
    para(doc, 'Probestunde — 60 €')
    para(doc, 'Einzelstunde — 70 €')
    para(doc, 'Duo-Training* (pro Person) — 40 €')
    para(doc, '*Preis pro Person bei 2:1-Training')
    para(doc, '[Button: Jetzt anfragen]')

    para(doc, 'Alle Preise inkl. MwSt. Mitgliedschaften und Karten auf persönliche Anfrage.')

    # ── Kontakt ───────────────────────────────────────────────────────────────
    h2(doc, 'Kontakt & Anfahrt')
    label(doc, 'Komm vorbei')

    para(doc, 'Studio Rüttenscheid', bold=True)
    para(doc, 'Gudulastraße 5 (Innenhof), 45131 Essen')
    para(doc, 'Pilates Matte · Barre Workout')

    para(doc, 'Studio Südviertel', bold=True)
    para(doc, 'Moltkestraße 16, 45128 Essen')
    para(doc, 'Reformer Pilates · Raumvermietung')

    para(doc, 'Telefon: 0201 857 721 37')
    para(doc, 'E-Mail: info@core-form.de')

    h3(doc, 'Kontaktformular')
    para(doc, 'Felder: Name · E-Mail · Telefon (optional) · Interesse (Dropdown) · Nachricht')
    para(doc, 'Interesse-Optionen: Reformer Pilates · Pilates Matte · Barre Workout · Personal Training · Raumvermietung · Ausbildung · Telefonischer Rückruf · Allgemeine Anfrage')
    para(doc, '[Button: Nachricht senden]')
    para(doc, 'Erfolg: Danke für deine Nachricht. Wir melden uns schnellstmöglich bei dir.')

    # ── Footer ────────────────────────────────────────────────────────────────
    h2(doc, 'Footer')
    para(doc, '© 2026 core:form Pilates Studio GbR, Essen')
    para(doc, 'Links: Impressum | Datenschutz | AGB | FAQ | Privacy')


# ─────────────────────────────────────────────────────────────────────────────

def buchung_haupt(doc):
    h1(doc, 'Buchungsseite (buchung.html)')
    divider(doc)

    label(doc, 'Studios & Kursangebote')
    h2(doc, 'Zwei Studios — eine Haltung.')
    para(doc, 'Bewegung braucht Aufmerksamkeit, Klarheit und Zeit. Wähle dein Studio – und finde das passende Format für dich.')

    para(doc, 'Matten- & Barre-Trainings finden im Studio Rüttenscheid statt, Reformer-Trainings im Studio Südviertel. Beide Studios verbindet derselbe Anspruch: präzise angeleitet, persönlich begleitet, in kleinen Gruppen.')
    para(doc, 'Buche direkt online im jeweiligen Studio – oder weiter unten dein Workshop-, Personal- oder Einsteigerangebot.')

    h3(doc, 'Studio 01 — Rüttenscheid')
    para(doc, 'Pilates Matte & Barre Workout')
    para(doc, '[Button: Info & Buchung]')

    h3(doc, 'Studio 02 — Südviertel')
    para(doc, 'Reformer Pilates & Raumvermietung')
    para(doc, '[Button: Info & Buchung]')

    label(doc, 'Online buchen')
    h2(doc, 'Geschlossene Kurse, Workshops & Ausbildung.')
    para(doc, 'Hier findest du alle übergreifenden Angebote von core:form: Einsteigerblöcke, Ausbildungsmodule und Workshops. Für laufende Stunden im jeweiligen Studio nutze die Buttons oben.')
    para(doc, '[Eversports-Buchungswidget — alle Studios + Workshops + Ausbildung]')


def buchung_rue(doc):
    h1(doc, 'Buchung Studio Rüttenscheid (buchung-ruettenscheid.html)')
    divider(doc)

    label(doc, 'Studio 01 · Rüttenscheid')
    h2(doc, 'Pilates & Barre im Herzen von Rüttenscheid.')
    para(doc, 'Persönlich, klar und langjährig erfahren. Mattentraining und Barre Workout in der Gudulastraße 5 (Innenhof) – in Kleingruppen und mit Aufmerksamkeit für jede Bewegung.')

    para(doc, 'Pilates Matte – Mattentraining in Kleingruppen (bis 12 Teilnehmer:innen) mit Fokus auf Bewegungsqualität, Atmung und Körperwahrnehmung.', bold=False)
    para(doc, 'Barre Workout – Kombination aus Pilates, Tanz, Kraft- und Faszientraining. Dynamisch und intensiv.')
    para(doc, 'Für Reformer-Trainings wechsle bitte zum Studio Südviertel.')

    label(doc, 'Online buchen')
    h2(doc, 'Termin im Studio Rüttenscheid.')
    para(doc, 'Das Buchungstool ist auf Matten- und Barre-Trainings im Studio Rüttenscheid (Gudulastraße 5, Innenhof) vorgefiltert. Wähle Format, Trainer:in und Zeitfenster.')
    para(doc, '[Eversports-Buchungswidget — Studio Rüttenscheid]')


def buchung_sued(doc):
    h1(doc, 'Buchung Studio Südviertel (buchung-suedviertel.html)')
    divider(doc)

    label(doc, 'Studio 02 · Südviertel')
    h2(doc, 'Reformer Pilates im Südviertel.')
    para(doc, 'Gruppentraining am Reformer für bis zu 9 Teilnehmer:innen – in der Moltkestraße 16. Strukturiert, technikorientiert und mit echter Wirkung.')

    para(doc, 'Reformer Essentials – strukturiertes, technikorientiertes Training mit Fokus auf Atmung und Ausrichtung. Ideal für Anfänger:innen und für alle, die ihre Basis stärken möchten.')
    para(doc, 'Reformer Flow – dynamisches Training für Teilnehmer:innen mit Pilates-Erfahrung. Fließend, kraftvoll, präzise.')
    para(doc, 'Reformer Gentle Flow – sanftes, moderates Training, donnerstags um 11:00 Uhr.')
    para(doc, 'Zusätzlich: Einsteigerkurse (6 Wochen), Reformer-Intro (75 Minuten), Personal Training und Workshops.')
    para(doc, 'Für Matte- & Barre-Trainings wechsle bitte zum Studio Rüttenscheid.')

    label(doc, 'Online buchen')
    h2(doc, 'Termin im Studio Südviertel.')
    para(doc, 'Das Buchungstool ist auf Reformer-Trainings im Studio Südviertel (Moltkestraße 16, 45128 Essen) voreingestellt. Filtere nach Format, Trainer:in oder Zeitfenster.')
    para(doc, '[Eversports-Buchungswidget — Studio Südviertel]')


# ─────────────────────────────────────────────────────────────────────────────

def ausbildung(doc):
    h1(doc, 'Ausbildungsseite (ausbildung.html)')
    divider(doc)

    # Hero
    label(doc, 'Ausbildung 2026')
    h2(doc, 'Reformer-Lehrerausbildung.')
    para(doc, 'Umfassende Ausbildung für angehende Pilates-Reformer-Trainer:innen in Essen. Fundiert, praxisnah und modern. Strukturiertes, intensives Lernen in wertschätzender Atmosphäre – mit einem erfahrenen Ausbildungsteam an deiner Seite.')
    para(doc, '[Button: Stimmen zur Ausbildung]  [Button: Termine ansehen]')

    # USPs
    label(doc, 'Was uns auszeichnet')
    h2(doc, 'Das Besondere an core:form.')

    usps = [
        ('01', 'Praxisnah ab Tag 1', 'Vom ersten Termin an arbeiten wir direkt am Reformer – nicht in der Theorie, sondern am Gerät.'),
        ('02', 'Persönliche Begleitung', 'Individuelle Aufmerksamkeit für jede Teilnehmer:in. Wir sehen, wo du stehst, und holen dich dort ab.'),
        ('03', 'Klare Didaktik', 'Komplexe Inhalte verständlich vermittelt – Schritt für Schritt, ohne Umweg über Fachjargon.'),
        ('04', 'Realistische Vorbereitung', 'Du wirst nicht für ein Lehrbuch ausgebildet, sondern für reale Gruppenkurse mit echten Menschen.'),
        ('05', 'Wertschätzende Lernkultur', 'Eine motivierende Atmosphäre, in der Fragen erwünscht sind und Fehler dazugehören.'),
        ('06', 'Erfahrenes Team', 'Eva mit über 25 Jahren Pilates-Erfahrung, Jelena mit über 10 Jahren Reformer-Praxis und einem Tanz-Background.'),
        ('07', 'Lernen im echten Studiobetrieb', 'Du erlebst von Anfang an, wie ein professionelles Studio funktioniert – nicht im sterilen Schulungsraum.'),
        ('08', 'Hochwertiges Skript', 'Strukturiertes Ausbildungsskript zum Mitnehmen – auch nach der Ausbildung dein Nachschlagewerk.'),
        ('09', 'Austausch & Wachstum', 'Feedback in der Gruppe, gegenseitiges Unterrichten, gemeinsam besser werden.'),
    ]
    for nr, titel, desc in usps:
        p = doc.add_paragraph()
        p.add_run(f'{nr}  {titel}').bold = True
        para(doc, desc)

    # Curriculum
    label(doc, 'Curriculum')
    h2(doc, 'Was du lernen wirst.')
    para(doc, 'Die Ausbildung deckt das vollständige Handwerkszeug einer Reformer-Trainer:in ab – vom Übungskanon über Cueing bis zur Arbeit mit unterschiedlichen Zielgruppen.')

    label(doc, 'Wichtiger Hinweis')
    para(doc, 'Die Ausbildung ist ein studiointernes Qualifizierungsformat und keine verbandlich zertifizierte Ausbildung. Nach Abschluss kannst du Personaltrainings und/oder Gruppenkurse anleiten.')

    inhalte = [
        ('Ca. 80 Reformer-Grundübungen', 'Mit zahlreichen Varianten und Modifikationen für jedes Niveau.'),
        ('Anatomisch sinnvolle Bewegungsabläufe', 'Verstehen, warum eine Übung wirkt – nicht nur, wie sie aussieht.'),
        ('Aufbau strukturierter Reformer-Stunden', 'Stundenbilder konstruieren, die didaktisch und anatomisch aufgehen.'),
        ('Teaching-Tipps & Cueing-Techniken', 'Verbal anleiten, sodass deine Teilnehmer:innen es spüren.'),
        ('Hands-on-Techniken', 'Sicheres taktiles Korrigieren mit Respekt und Klarheit.'),
        ('Sicherheitsprinzipien', 'Geräteumgang, Federwiderstand, häufige Risikomomente.'),
        ('Übungen für Einsteiger:innen, Flow-Kurse & Mixed Levels', 'Programme für unterschiedliche Erfahrungsstufen entwickeln.'),
        ('Umgang mit verschiedenen Zielgruppen', 'Anfänger:innen, Fortgeschrittene, ältere Teilnehmer:innen, Männer.'),
        ('Beschwerden & Kontraindikationen', 'Wann eine Übung nicht angezeigt ist und welche Alternativen es gibt.'),
    ]
    for titel, desc in inhalte:
        p = doc.add_paragraph()
        p.add_run('✓  ' + titel).bold = True
        para(doc, '   ' + desc)

    # Termine
    label(doc, 'Aufbau & Termine')
    h2(doc, 'Vier Präsenztage, begleitet von eigener Praxis.')

    para(doc, 'Präsenztermine — jeweils 10:00 bis 16:00 Uhr', bold=True)
    para(doc, 'Modul 01 — 25.04.2026')
    para(doc, 'Modul 02 — 16.05.2026')
    para(doc, 'Modul 03 — 20.06.2026')
    para(doc, 'Modul 04 — 11.07.2026')
    para(doc, 'Ort: core:form, Moltkestraße 16, Essener Südviertel')

    para(doc, 'Begleitende Praxis (Pflichtanteil)', bold=True)
    para(doc, 'Eigene Praxis: Mindestens 25 Stunden Reformer-Training – idealerweise bei verschiedenen Trainer:innen und in verschiedenen Studios.')
    para(doc, 'Stundenprotokolle: Mindestens 5 dokumentierte Einheiten, um die Perspektive als Unterrichtende:r zu entwickeln.')
    para(doc, 'Hospitationen: Optional und empfohlen: Möglichkeit, in laufenden core:form-Kursen zu assistieren.')

    # Zielgruppe
    label(doc, 'Für wen?')
    h2(doc, 'Diese Ausbildung ist für dich.')
    para(doc, 'Wenn du Pilates ernst nimmst und Reformer professionell unterrichten möchtest, bist du hier richtig. Auch wenn du bereits in verwandten Bereichen arbeitest, profitierst du von der Tiefe und Struktur unserer Ausbildung.')
    zielgruppe = [
        'Du bist begeistert von Pilates und möchtest tiefer einsteigen.',
        'Du willst Reformer professionell unterrichten – als Hauptberuf oder zweites Standbein.',
        'Du arbeitest bereits als Trainer:in im Bereich Fitness, Yoga oder Tanz und möchtest dein Repertoire erweitern.',
        'Du willst deine eigene Technik vertiefen und Kenntnisse fundieren.',
        'Du möchtest später bei core:form oder in anderen Studios Gruppen oder Personaltraining unterrichten.',
    ]
    for z in zielgruppe:
        para(doc, '• ' + z)

    label(doc, 'Voraussetzung')
    para(doc, 'Wer keine Pilates-Mattenausbildung absolviert hat, sollte am Pilates-Grundlagen-Kompaktseminar am 18.04.2026 teilnehmen (zusätzlich 250 €). Hier werden die wesentlichen Pilates-Techniken sowie die zugrundeliegenden Prinzipien vermittelt.')

    # Team
    label(doc, 'Ausbildungsteam')
    h2(doc, 'Eva & Jelena — im Team-Teaching.')

    para(doc, 'Ausbildungsleitung & Inhaberin', bold=True)
    para(doc, 'Eva Wiedwald', bold=True)
    para(doc, 'Eva unterrichtet seit über 25 Jahren Pilates, davon 14 Jahre hauptberuflich. Ihr Unterricht ist geprägt von Klarheit und einer tiefen Kenntnis des Pilates-Systems.')
    para(doc, 'Durch langjährige Erfahrung im Reformer- und Mattentraining hat sie einen besonderen Blick dafür entwickelt, wie Menschen sinnvoll, sicher und mit Freude bewegt werden können. Ihre Stärke ist es, komplexe Inhalte verständlich zu machen und jeder Person genau dort zu begegnen, wo sie steht.')
    para(doc, 'Sie leitet die Ausbildung strukturiert, verständlich und mit viel Herz.')

    para(doc, 'Ausbildungsassistenz', bold=True)
    para(doc, 'Jelena Grjasnowa', bold=True)
    para(doc, 'Jelena war viele Jahre professionelle Balletttänzerin und unterrichtet seit über 10 Jahren Reformertraining – im Personaltraining wie auch in der Gruppe.')
    para(doc, 'Sie bringt Dynamik, Präzision und einen wunderbaren Blick für Technik und Bewegungsqualität mit. In der Ausbildung unterstützt sie besonders bei Technik, Flow und der sicheren Ausführung am Gerät.')

    para(doc, 'Auszeichnungen: Premium-Mitglied Deutscher Pilates Verband · Aktive Mitgliedschaft 2026')

    # Investition
    label(doc, 'Investition')
    h2(doc, 'Deine Investition in echtes Können.')

    para(doc, 'Reformer-Lehrerausbildung 2026', bold=True)
    para(doc, '1.390 € inkl. MwSt.')
    para(doc, '• Alle 4 Präsenztermine (jeweils 10–16 Uhr)')
    para(doc, '• Ausführliches Ausbildungsskript')
    para(doc, '• Teilnahmebescheinigung nach Abschluss')
    para(doc, '• Individuelle Betreuung durch Eva & Jelena')
    para(doc, '• Nutzung der Reformer zum Selbststudium (nach Verfügbarkeit)')
    para(doc, 'Ratenzahlung möglich.')

    para(doc, 'Optional: Pilates-Grundlagen-Seminar — + 250 €', bold=True)
    para(doc, 'Kompaktseminar am 18.04.2026 für alle ohne abgeschlossene Mattenausbildung. Vermittelt die wesentlichen Pilates-Techniken und Prinzipien.')

    para(doc, 'Optional: Practice Day — + 120 €', bold=True)
    para(doc, 'Zusätzlicher Praxistag zur Vertiefung. Ideal als Auffrischung kurz vor dem ersten eigenen Unterrichten.')

    # Bewerbung
    label(doc, 'Bewerbung & Erstgespräch')
    h2(doc, 'Wir lernen uns vorher persönlich kennen.')
    para(doc, 'Vor jeder Anmeldung steht ein persönliches Erstgespräch. Wir möchten sicher gehen, dass die Ausbildung für dich passt – und du bekommst ein klares Bild davon, was dich erwartet. Das Gespräch ist unverbindlich und kostenlos.')
    bewerbung = [
        'Deine Motivation & Ziele',
        'Deine Vorerfahrung mit Pilates und Bewegung',
        'Ob die Ausbildung aktuell für dich geeignet ist',
        'Ob du das Grundlagen-Seminar brauchst',
        'Alle offenen Fragen',
    ]
    for b in bewerbung:
        para(doc, '• ' + b)

    # Testimonials
    label(doc, 'Stimmen aus der Ausbildung')
    h2(doc, 'Was Trainer:innen über uns sagen.')

    testimonials = [
        ('Elle', 'Trainerin',
         'Die Ausbildung bei Jelena und Eva ist besonders toll, weil sie durch hohe fachliche Qualität, klare Struktur und eine außergewöhnlich wertschätzende Lernatmosphäre überzeugt. Weil man hier nicht nur fachlich geschult wird, sondern auch als Mensch wirklich gesehen wird. Weil die beiden eine Umgebung schaffen, in der man sich sicher, unterstützt und motiviert fühlt. Und weil ihre Kombination aus Kompetenz, Authentizität, Herz und feinem Humor jede Einheit besonders macht und fachliches wie persönliches Wachstum ermöglicht. Eine Ausbildung, die nachhaltig wirkt und für die ich sehr dankbar bin.'),
        ('Franziska', 'Trainerin',
         'Eva begleitet einen mit ihrer wertschätzenden, offenen und humorvollen Art durch die Reformer-Ausbildung. Ich habe von ihrem hohen Fachwissen als Pilates-Trainerin absolut profitiert. Eine super Ergänzung ist das Ausbildungs-Skript, welches nicht nur beim Lernen zwischen den praktischen Tagen unterstützt, sondern mit viel Herzblut und Fachwissen von Eva gespickt ist. Man merkt, wie wichtig ihr ist, den Auszubildenden nicht nur die Übungen, sondern auch Fachwissen und didaktische Grundlagen beizubringen, damit ich als Trainerin sicher durchstarten kann. Danke dafür, Eva.'),
        ('Marina', 'Trainerin',
         'Die Ausbildung zur Reformer-Trainerin bei Eva von core:form ist für mich eine sehr bereichernde Erfahrung. Mit ihrer offenen, wertschätzenden Art leitet Eva sowohl die Ausbildung als auch ihre Kurse. Ihr Fachwissen und ihre Erfahrung werden in jedem Modul der Ausbildung deutlich. Besonders schätze ich die Praxisnähe, das aufwändig ausgearbeitete Handout und die klar strukturierte Vermittlung der Inhalte. Schritt für Schritt begleitet Eva den Weg zur Trainerin – mit fundiertem Hintergrundwissen und einer individuellen Vorbereitung auf die eigene Trainerpraxis. Auch im Studioalltag zeigen sich die hohe Qualität der Trainer:innen und die persönliche, wertschätzende Begleitung.'),
        ('Rolf', 'Trainer',
         'Seit über 13 Jahren trainiere ich Pilates am Reformer unter Evas Anleitung. Dadurch hat sich mein Fitnessniveau deutlich verbessert, und das Training ist zu einem festen Bestandteil meiner Gesundheit geworden. Evas außergewöhnliche Präzision, ihr geschulter Blick und ihre Fähigkeit, Bewegungen differenziert zu analysieren und gezielt zu korrigieren, haben mich über all die Jahre nachhaltig geprägt. Der konsequente Höhepunkt dieser langjährigen Praxis war für mich die Ausbildung zum Trainer, die ich bei Eva und Jelena im Team-Teaching absolvieren durfte. Diese Zusammenarbeit ist mir als besonders stimmiges Gesamtbild in Erinnerung geblieben: Eva mit ihrer klaren Struktur und der konsequenten Vermittlung der Pilates-Prinzipien, und Jelena, die mit großer Präzision in der Ausführung und Vermittlung der Übungen meine Ausbildung auf ideale Weise ergänzt hat. Ich bin sehr dankbar, von beiden gelernt zu haben.'),
        ('Selina', 'Trainerin',
         'Die Ausbildung bei Eva ist total praxisnah aufgebaut, da man in jeder Einheit spürt, dass Eva aus 20 Jahren Erfahrung spricht und dieses Wissen direkt anwendbar macht. Zusammen mit Jelena ergibt das eine super Kombi, die die Inhalte nicht nur verständlich, sondern auch inspirierend vermittelt.'),
        ('Shari', 'Trainerin',
         'Vom ersten Termin an habe ich mich unglaublich gut aufgehoben gefühlt. An der Ausbildung bei Eva und Jelena hat mir besonders die persönliche Betreuung gefallen. Lernen macht hier nicht nur Spaß, sondern motiviert richtig.'),
    ]
    for name, rolle, text in testimonials:
        p = doc.add_paragraph()
        p.add_run(f'{name} — {rolle}').bold = True
        para(doc, text)
        doc.add_paragraph()

    # Final CTA
    label(doc, 'Reformer-Lehrerausbildung 2026')
    h2(doc, 'Bereit, deine Praxis zur Profession zu machen?')
    para(doc, 'Plätze sind begrenzt. Schreib uns für ein unverbindliches Erstgespräch oder ruf direkt im Studio an.')
    para(doc, '[Button: Erstgespräch anfragen]  [Button: 0201 857 721 37]')


# ─────────────────────────────────────────────────────────────────────────────

def galerie_seite(doc):
    h1(doc, 'Galerie (galerie.html)')
    divider(doc)

    label(doc, 'Studio & Atmosphäre')
    h2(doc, 'Bewegung, festgehalten.')
    para(doc, 'Ein Blick in unsere Räume in Rüttenscheid und im Südviertel — Trainingsmomente, Reformer, Matte und das, was core:form ausmacht.')

    h3(doc, 'Bilder in der Galerie')
    bilder = [
        ('1',  'COREFORM_web_001.jpg', 'Studio-Eindruck 1'),
        ('2',  'COREFORM_web_002.jpg', 'Studio-Eindruck 2'),
        ('3',  'COREFORM_web_003.jpg', 'Studio-Eindruck 3'),
        ('4',  'COREFORM_web_004.jpg', 'Studio-Eindruck 4'),
        ('5',  'COREFORM_web_005.jpg', 'Studio-Eindruck 5'),
        ('6',  'COREFORM_web_006.jpg', 'Studio-Eindruck 6'),
        ('7',  'COREFORM_web_007.jpg', 'Studio-Eindruck 7'),
        ('8',  'COREFORM_web_008.jpg', 'Studio-Eindruck 8'),
        ('9',  'COREFORM_web_009.jpg', 'Studio-Eindruck 9'),
        ('10', 'COREFORM_web_010.jpg', 'Studio-Eindruck 10'),
        ('11', 'COREFORM_web_011.jpg', 'Studio-Eindruck 11'),
        ('12', 'COREFORM_web_012.jpg', 'Studio-Eindruck 12'),
        ('13', 'COREFORM_web_013.jpg', 'Studio-Eindruck 13'),
        ('14', 'COREFORM_web_014.jpg', 'Studio-Eindruck 14'),
        ('15', 'COREFORM_web_015.jpg', 'Studio-Eindruck 15'),
        ('16', 'COREFORM_web_016.jpg', 'Studio-Eindruck 16'),
        ('17', 'Eva-Pilates 02.jpg',   'Eva Pilates 2'),
        ('18', 'Eva-Pilates 03.jpg',   'Eva Pilates 3'),
        ('19', 'Eva-Pilates 04.jpg',   'Eva Pilates 4'),
        ('20', 'Eva-Pilates 05.jpg',   'Eva Pilates 5'),
        ('21', 'Eva-Pilates 06.jpg',   'Eva Pilates 6'),
        ('22', 'Eva-Pilates 07.jpg',   'Eva Pilates 7'),
        ('23', 'Eva-Pilates 08.jpg',   'Eva Pilates 8'),
        ('24', 'Eva-Pilates 09.jpg',   'Eva Pilates 9'),
        ('25', 'Eva-Pilates 10.jpg',   'Eva Pilates 10'),
        ('26', 'Eva-Pilates 11.jpg',   'Eva Pilates 11'),
    ]
    for nr, datei, alt in bilder:
        para(doc, f'{nr}. {datei} — {alt}')


# ─────────────────────────────────────────────────────────────────────────────

def videos_seite(doc):
    h1(doc, 'Videos (videos.html)')
    divider(doc)

    label(doc, 'Unser Trainingsangebot')
    h2(doc, 'Bewegung in Bewegung.')
    para(doc, 'Du suchst nach einem Training, das nicht nur effektiv ist, sondern dir auch ein gutes Körpergefühl schenkt? Unsere Angebote verbinden Bewegung, Achtsamkeit und Freude – für jedes Alter und jedes Level. Ob mit dem eigenen Körpergewicht oder an modernen Geräten: Unser Ziel ist es, dich in deiner Balance, Kraft und Beweglichkeit zu unterstützen. Klick auf ein Video, um mehr über das Format zu erfahren.')

    kacheln = [
        ('Pilates Matte', 'Matte', 'Matte Video.mp4',
         'Das klassische Pilates-Training auf der Matte bildet die Grundlage für eine stabile Körpermitte und eine gesunde Haltung. Durch kontrollierte Bewegungen wird die Tiefenmuskulatur aktiviert und die Körperwahrnehmung geschult.\n'
         'Ideal für alle, die ein ganzheitliches Training in ruhiger Atmosphäre suchen.'),

        ('Reformer Pilates', 'Reformer', 'Reformer Boom Bang.mp4',
         'Reformer-Pilates ist so beliebt wie nie und das nicht ohne Grund: Die kontrollierten, fließenden Bewegungen auf dem Reformer trainieren die Tiefenmuskulatur bzw. Muskelketten, gleichzeitig wird die Körperwahrnehmung geschult, muskuläre Dysbalancen ausgeglichen und die Körperhaltung verbessert.\n'
         'Der Widerstand lässt sich durch das Federsystem individuell anpassen, sodass das Training sehr effektiv ist. Intensives und zugleich gelenkschonendes Training auf dem Pilates-Reformer.\n'
         'Die geführten Bewegungen ermöglichen ein besonders präzises Arbeiten — perfekt für gezielten Muskelaufbau, Rehabilitation und das Erlernen sauberer Technik.'),

        ('Reformer Pilates', 'Reformer Gruppe', 'Reformer Harder.mp4',
         'Wenn der Reformer dir vertraut ist und du nach mehr Intensität suchst — dann verschiebe gerne die Schwelle. Höhere Federwiderstände, dynamischere Übergänge und komplexere Bewegungsmuster, immer noch mit der Präzision, die Pilates ausmacht.\n'
         'Voraussetzung: solide Grundlagen aus Einsteiger- und Mittelstufenkursen. Eine Stunde, die fordert — und die Belohnung ist spürbare Tiefenkraft.'),

        ('Barre Workout', 'Barre Harmonisch', 'Barre Harmony.mp4',
         'Barre Workout ist eine spannende Mischung aus Ballett, funktionellen Übungen und Yoga, welche den ganzen Körper sanft stärkt und formt. Wir arbeiten mit und auf Musik. Je nach Trainingsablauf kommt Equipment wie Bälle, Bänder oder kleine Gewichte gezielt zum Einsatz.\n'
         'Barre Workout ist für dich geeignet, wenn du gesund und nicht schwanger bist. Bitte bring zum Training rutschfeste Socken, ein Handtuch und Wasser mit. Alles andere stellen wir dir zur Verfügung.\n'
         'Die 60-minütige Trainingseinheit ist ein offener Kurs, in den du jederzeit einsteigen kannst. Keine Tanzerfahrung erforderlich.'),

        ('Barre Workout', 'Barre Intensiv', 'Barre Burn.mp4',
         'Barre geht auch intensiver als im klassischen Barre Workout. Schnellere Sequenzen, mehr Wiederholungen, gezielter Pulse-Burnout an Bauch, Beinen und Po.\n'
         'Musik treibt, Atmung führt — und am Ende glüht die ganze Muskulatur. Für alle, die schon Erfahrung mit Barre haben und ihre Limits gerne testen.'),

        ('Pilates & Lebensgefühl', 'Cihangir, Istanbul', 'Cihangir Pilates.mp4',
         'Pilates ist mehr als Training — es ist ein Lebensgefühl. Dieses Video entstand in Istanbul, im lebhaften Stadtteil Cihangir: Eva am Reformer und am Cadillac, und danach ein Streifzug durch eine der faszinierendsten Städte der Welt.\n'
         'Was Joseph Pilates einst in Mönchengladbach begann und nach New York trug, ist heute überall angekommen. In jedem gut geführten Studio, auf jedem Kontinent. Die Methode hat die Welt erobert — weil sie funktioniert.'),

        ('Neugier & Methode', 'Reformer, Istanbul', 'Pilates Istanbul.mp4',
         'Kein Studio gleicht dem anderen — und das ist eine gute Sache. Während eines Aufenthalts in Istanbul trainierte Eva an einem Reformer mit einer ihr unbekannten Programmfolge: andere Akzente, neue Übergänge, ein frischer Blick auf die vertraute Methode.\n'
         'Für Eva ist Pilates kein abgeschlossenes Wissen, sondern ein ständiger Wegbegleiter. Es gibt immer etwas zu entdecken, immer jemanden, von dem man lernen kann. Diese Neugier ist Teil der Haltung, die core:form ausmacht.'),
    ]

    for rubrik, titel, datei, beschreibung in kacheln:
        label(doc, rubrik)
        h3(doc, titel)
        para(doc, f'[Video: {datei}]')
        for zeile in beschreibung.split('\n'):
            if zeile.strip():
                para(doc, zeile)
        doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────

def faq_seite(doc):
    h1(doc, 'FAQ (faq.html)')
    divider(doc)

    label(doc, 'Häufige Fragen')
    h2(doc, 'FAQ.')
    para(doc, 'Antworten auf die Fragen, die uns am häufigsten erreichen. Du findest deine Antwort nicht? Schreib uns einfach.')

    fragen = [
        ('Habt ihr eigentlich nur Mattenkurse, oder bietet ihr mehr?',
         'Matte, Reformer, Barre, TRX und vieles mehr.'),
        ('Bietet ihr auch Einzeltraining?',
         'Ja. Auch am Reformer und sogar für zwei Personen gleichzeitig als Duo-Training.'),
        ('Warum habt ihr zwei Standorte?',
         'Begonnen hat alles in der Gudulastr. ohne Reformer Geräte. Später kamen am selben Standort zwei Reformer mit Tower Aufsatz für Einzel- oder Zweiertraining hinzu. Das Training am Reformer liegt inzwischen aber derart im Trend und die Nachfrage wurde entsprechend hoch, dass wir uns entschieden haben ein Gruppentraining für mehr als zwei anzubieten. Dafür brauchten wir neue Räume, die wir seit dem 01. November 2025 in der Moltkestr. 16 eingerichtet haben für bis zu 9 Teilnehmer gleichzeitig.'),
    ]

    for frage, antwort in fragen:
        para(doc, frage, bold=True)
        para(doc, antwort)
        doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────

def newsletter(doc):
    h1(doc, 'Newsletter-Popup (data/newsletter.html)')
    divider(doc)

    h2(doc, 'Neues aus dem Studio.')

    para(doc, 'Liebe core:form Community,')
    para(doc, 'für den Mai haben sich noch einige Änderungen und zusätzliche Termine ergeben – deshalb hier ein kurzes Update für Dich.')

    h3(doc, 'Feiertage im Mai')
    para(doc, 'An folgenden Tagen findet kein regulärer Unterricht statt:')
    para(doc, '• Freitag, 1. Mai – Tag der Arbeit')
    para(doc, '• Donnerstag, 14. Mai – Christi Himmelfahrt')
    para(doc, '• Montag, 25. Mai – Pfingstmontag')
    para(doc, 'Wenn Du von einem Ausfall betroffen bist, weiche bitte gern auf einen anderen Termin aus.')

    h3(doc, 'Termine & Änderungen im Überblick')

    termine = [
        ('Freitag, 8. Mai · 11:00 Uhr',
         'Reformer Flow mit PIA findet doch statt (zuvor als Ausfall angekündigt)'),
        ('Samstag, 9. Mai · 12:15 – 13:30 Uhr',
         'Reformer Intro mit JELENA'),
        ('Donnerstag, 14. Mai · 18:00 Uhr (Feiertag)',
         'Barre Workout mit CRISTINA'),
        ('Freitag, 15. Mai · 17:30 Uhr',
         'Charity-Stunde: HIIT Pilates mit JULIA · 75 Minuten\n→ auf Spendenbasis, Einnahmen für den guten Zweck'),
        ('Donnerstag, 21. Mai · 19:30 Uhr',
         'Start neuer Reformer-Einsteigerkurs mit KATJA'),
        ('Samstag, 23. Mai · 12:15 – 13:30 Uhr',
         'Reformer Intro mit JELENA'),
        ('Samstag, 30. Mai – Special mit CLAUDIA LAU',
         '11:15 Uhr – Reformer (Intermediate Level)\n12:30 Uhr – Reformer Jumpboard\nDie Buchung erfolgt direkt über mich (nicht über Eversports).'),
        ('Samstag, 31. Mai · 10:30 – 12:00 Uhr',
         'Refine your Reformer – Technik-Workshop mit JELENA\nStudio Rüttenscheid\nEin Workshop auf der Matte, besonders für alle, die regelmäßig am Reformer trainieren und ihre Pilates-Basics, Ausrichtung und Technik noch einmal gezielt vertiefen möchten.'),
    ]

    for datum, inhalt in termine:
        para(doc, datum, bold=True)
        for zeile in inhalt.split('\n'):
            if zeile.strip():
                para(doc, zeile)
        doc.add_paragraph()

    h3(doc, 'Hinweis zu Mitgliedschaften & Karten')
    para(doc, 'Mitgliedschaften sowie ermäßigte Mitgliedschaften und ermäßigte 10er-Karten sind ausschließlich nach Rücksprache mit uns erhältlich. Sprich uns dazu einfach gern direkt an.')

    para(doc, 'Wenn Du Fragen hast oder unsicher bist, welcher Termin gut für Dich passt, melde Dich jederzeit gern bei mir. Ich freue mich, Dich im Mai im Studio zu sehen. Herzliche Grüße und einen guten Start in den Tag')
    para(doc, 'Bis bald im Studio, deine Eva & das core:form Team')


# ─────────────────────────────────────────────────────────────────────────────

def seo_sektion(doc):
    h1(doc, 'SEO-Inhalte')
    divider(doc)

    # Startseite
    h2(doc, 'Startseite (index.html)')
    seo_startseite = [
        ('Title', 'core:form — Pilates & Reformer Studio Essen | Rüttenscheid & Südviertel'),
        ('Meta Description', 'Pilates Matte, Reformer Pilates und Barre Workout in Essen. Zwei Studios in Rüttenscheid und Südviertel — kleine Gruppen, durchdachte Programme, klare Haltung.'),
        ('Keywords', 'Pilates Essen, Reformer Pilates Essen, Barre Workout Essen, Pilates Studio Rüttenscheid, Pilates Südviertel, Personal Training Essen, Pilates Ausbildung Essen'),
        ('Robots', 'index, follow, max-image-preview:large'),
        ('Canonical', 'https://core-form.de/'),
        ('OG Title', 'core:form — Pilates & Reformer Studio Essen'),
        ('OG Description', 'Pilates Matte, Reformer Pilates und Barre Workout in Essen. Zwei Studios in Rüttenscheid und Südviertel.'),
        ('OG Image', 'https://core-form.de/media/COREFORM_web_015.jpg'),
        ('Twitter Title', 'core:form — Pilates & Reformer Studio Essen'),
        ('Twitter Description', 'Pilates Matte, Reformer Pilates und Barre Workout in Essen. Zwei Studios in Rüttenscheid und Südviertel.'),
    ]
    for feld, wert in seo_startseite:
        p = doc.add_paragraph()
        p.add_run(f'{feld}: ').bold = True
        p.add_run(wert)

    h3(doc, 'Strukturierte Daten — Organization (JSON-LD)')
    para(doc, 'Name: core:form Pilates Studio')
    para(doc, 'URL: https://core-form.de/')
    para(doc, 'E-Mail: info@core-form.de')
    para(doc, 'Telefon: +49 201 857 721 37')
    para(doc, 'Standort 1: core:form Studio Rüttenscheid — Gudulastraße 5, 45131 Essen — Pilates Matte und Barre Workout')
    para(doc, 'Standort 2: core:form Studio Südviertel — Moltkestraße 16, 45128 Essen — Reformer Pilates und Raumvermietung')
    para(doc, 'Social: instagram.com/coreform_essen · facebook.com/evapilates.de · youtube.com/channel/UCPZpD1CyB2Vo82C0-_lxUMg')

    h3(doc, 'Strukturierte Daten — WebSite (JSON-LD)')
    para(doc, 'Name: core:form')
    para(doc, 'URL: https://core-form.de/')
    para(doc, 'Sprache: de-DE')

    doc.add_paragraph()

    # Ausbildung
    h2(doc, 'Ausbildungsseite (ausbildung.html)')
    seo_ausbildung = [
        ('Title', 'Reformer-Lehrerausbildung 2026 in Essen — core:form'),
        ('Meta Description', 'Reformer-Lehrerausbildung 2026 in Essen — praxisorientiertes Qualifizierungsformat für angehende Reformer-Trainer:innen mit erfahrenen Ausbilder:innen.'),
        ('Keywords', 'Reformer Ausbildung Essen, Pilates Ausbildung NRW, Reformer Trainer Ausbildung, Pilates Lehrerausbildung Essen'),
        ('Robots', 'index, follow, max-image-preview:large'),
        ('Canonical', 'https://core-form.de/ausbildung.html'),
        ('OG Title', 'Reformer-Lehrerausbildung 2026 — core:form Essen'),
        ('OG Description', 'Praxisorientierte Reformer-Lehrerausbildung 2026 in Essen. Kleine Gruppe, erfahrene Ausbilder:innen, durchdachte Module.'),
        ('OG Image', 'https://core-form.de/media/Eva-Pilates 01.jpg'),
        ('Twitter Title', 'Reformer-Lehrerausbildung 2026 — core:form Essen'),
        ('Twitter Description', 'Praxisorientierte Reformer-Lehrerausbildung 2026 in Essen. Kleine Gruppe, erfahrene Ausbilder:innen.'),
    ]
    for feld, wert in seo_ausbildung:
        p = doc.add_paragraph()
        p.add_run(f'{feld}: ').bold = True
        p.add_run(wert)

    h3(doc, 'Strukturierte Daten — Course (JSON-LD)')
    para(doc, 'Name: Reformer-Lehrerausbildung 2026')
    para(doc, 'Beschreibung: Studiointernes Qualifizierungsformat für angehende Reformer-Pilates-Trainer:innen. Praxisorientierte Ausbildung mit erfahrenen Ausbilder:innen in kleiner Gruppe.')
    para(doc, 'Anbieter: core:form Pilates Studio')
    para(doc, 'Sprache: de-DE')
    para(doc, 'Kurstyp: blended (Präsenz + Eigenpraxis)')
    para(doc, 'Bildungsniveau: Berufliche Weiterbildung')
    para(doc, 'Ort: Moltkestraße 16, 45128 Essen')

    doc.add_paragraph()

    # Buchung
    h2(doc, 'Buchungsseite (buchung.html)')
    seo_buchung = [
        ('Title', 'Online buchen — Pilates, Reformer & Barre in Essen | core:form'),
        ('Meta Description', 'Online-Buchung core:form Essen: Pilates Matte, Reformer Pilates, Barre Workout — Studios Rüttenscheid und Südviertel. Termine, Workshops, Ausbildung.'),
        ('Keywords', 'Pilates buchen Essen, Reformer Pilates buchen, Barre Workout Essen buchen, Pilates Termin Essen'),
        ('Robots', 'index, follow, max-image-preview:large'),
        ('Canonical', 'https://core-form.de/buchung.html'),
        ('OG Title', 'Online buchen — core:form Essen'),
        ('OG Description', 'Pilates, Reformer und Barre Workout in Essen online buchen. Studios Rüttenscheid und Südviertel.'),
        ('OG Image', 'https://core-form.de/media/COREFORM_web_015.jpg'),
    ]
    for feld, wert in seo_buchung:
        p = doc.add_paragraph()
        p.add_run(f'{feld}: ').bold = True
        p.add_run(wert)


# ─────────────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    t = doc.add_heading('core:form — Website-Texte', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('Vollständige Textdokumentation aller Seiten (1:1 aus den HTML-Quelldateien)')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Stand: Mai 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    page(doc)

    startseite(doc);  page(doc)
    buchung_haupt(doc); page(doc)
    buchung_rue(doc);   page(doc)
    buchung_sued(doc);  page(doc)
    ausbildung(doc);    page(doc)
    galerie_seite(doc); page(doc)
    videos_seite(doc);  page(doc)
    faq_seite(doc);     page(doc)
    newsletter(doc);    page(doc)
    seo_sektion(doc)

    out = BASE / 'core-form_Website-Texte.docx'
    doc.save(out)
    print(f'Gespeichert: {out}')


if __name__ == '__main__':
    main()
